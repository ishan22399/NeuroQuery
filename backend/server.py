from fastapi import FastAPI, APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, ConfigDict
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timezone
import asyncio

# Document processing
import PyPDF2
import docx
from PIL import Image
import pytesseract
import io
import base64

# RAG components
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
from langchain_text_splitters import RecursiveCharacterTextSplitter

# LLM
from openai import AzureOpenAI

ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection
mongo_url = os.environ['MONGO_URL']
client = AsyncIOMotorClient(mongo_url)
db = client[os.environ['DB_NAME']]

# Initialize sentence transformer model (free)
embedding_model = SentenceTransformer('all-MiniLM-L6-v2')

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# FAISS index (in-memory)
faiss_index = None
chunk_metadata = {}  # Changed from list to dict for O(1) lookup
metadata_index_map = []  # Maps FAISS index position to chunk_id

async def rebuild_faiss_index():
    """Rebuild FAISS index from MongoDB on startup"""
    global faiss_index, chunk_metadata, metadata_index_map
    
    try:
        # Load all chunks from MongoDB
        chunks = await db.document_chunks.find({}, {"_id": 0}).to_list(10000)
        
        if not chunks:
            logging.info("No chunks found in MongoDB, FAISS index empty")
            faiss_index = None
            return
        
        # Extract embeddings and metadata
        embeddings_list = []
        chunk_metadata = {}  # Reset to empty dict
        metadata_index_map = []
        
        for chunk in chunks:
            if 'embedding' in chunk and chunk['embedding']:
                chunk_id = chunk.get('id')
                embeddings_list.append({
                    'embedding': np.array(chunk['embedding']),
                    'chunk_id': chunk_id,
                    'document_id': chunk.get('document_id'),
                    'document_name': chunk.get('document_name', ''),
                    'chunk_index': chunk.get('chunk_index', 0),
                    'text': chunk.get('text', '')
                })
        
        if not embeddings_list:
            logging.info("No embeddings found in MongoDB")
            faiss_index = None
            return
        
        # Initialize FAISS index with FlatL2 (simple, reliable, no training needed)
        embeddings_array = np.array([e['embedding'] for e in embeddings_list]).astype('float32')
        dimension = embeddings_array.shape[1]
        
        # Use simple FlatL2 index - reliable and works with any number of documents
        faiss_index = faiss.IndexFlatL2(dimension)
        faiss_index.add(embeddings_array)
        
        # Rebuild metadata as dictionary for O(1) lookup
        for idx, e in enumerate(embeddings_list):
            chunk_id = e['chunk_id']
            chunk_metadata[chunk_id] = {
                'document_id': e['document_id'],
                'document_name': e['document_name'],
                'chunk_index': e['chunk_index'],
                'text': e['text']
            }
            metadata_index_map.append(chunk_id)
        
        logging.info(f"FAISS index rebuilt from MongoDB: {len(chunk_metadata)} chunks loaded")
    except Exception as e:
        logging.error(f"Error rebuilding FAISS index: {e}")
        faiss_index = None

# Models
class Document(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    filename: str
    file_type: str
    upload_date: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    total_chunks: int = 0
    processed: bool = False

class DocumentChunk(BaseModel):
    model_config = ConfigDict(extra="ignore")
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    chunk_index: int
    text: str
    embedding: Optional[List[float]] = None

class QueryRequest(BaseModel):
    query: str
    mode: str = "detailed"  # concise, detailed, research
    document_ids: Optional[List[str]] = None

class Citation(BaseModel):
    chunk_id: str
    document_id: str
    document_name: str
    text: str
    similarity: float

class QueryResponse(BaseModel):
    answer: str
    citations: List[Citation]
    faithfulness_score: float
    retrieval_details: Optional[Dict[str, Any]] = None
    refused: bool = False

class ChatMessage(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    role: str  # "user" or "assistant"
    content: str
    citations: Optional[List[Citation]] = None
    timestamp: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

class ChatSession(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    title: str = "New Chat"
    messages: List[ChatMessage] = []
    document_ids: Optional[List[str]] = None
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))

# Helper functions
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text, links, and metadata from PDF"""
    try:
        pdf_file = io.BytesIO(file_bytes)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        # Extract metadata
        if pdf_reader.metadata:
            text += "=== DOCUMENT METADATA ===\n"
            if pdf_reader.metadata.title:
                text += f"Title: {pdf_reader.metadata.title}\n"
            if pdf_reader.metadata.author:
                text += f"Author: {pdf_reader.metadata.author}\n"
            if pdf_reader.metadata.subject:
                text += f"Subject: {pdf_reader.metadata.subject}\n"
            text += "\n"
        
        # Extract text and links from pages
        for page_num, page in enumerate(pdf_reader.pages, 1):
            text += f"\n--- Page {page_num} ---\n"
            
            # Extract text
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
            
            # Extract links/annotations
            if "/Annots" in page:
                text += "\n[Links on this page]\n"
                for annot in page["/Annots"]:
                    try:
                        obj = annot.get_object()
                        if obj["/Subtype"] == "/Link":
                            if "/A" in obj:
                                link_info = obj["/A"].get_object()
                                if "/URI" in link_info:
                                    uri = link_info["/URI"]
                                    text += f"- Link: {uri}\n"
                    except:
                        pass
        
        return text
    except Exception as e:
        logging.error(f"Error extracting PDF: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text, links, tables, and images from Word document"""
    try:
        doc_file = io.BytesIO(file_bytes)
        doc = docx.Document(doc_file)
        text = ""
        
        # Extract from paragraphs (preserves links)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
            
            # Extract hyperlinks from paragraph runs
            for run in para.runs:
                if run.element.rPr is not None:
                    rPr = run.element.rPr
                    if rPr.rStyle is not None:
                        # Check for hyperlinks
                        for child in run.element.iter():
                            if 'hyperlink' in child.tag.lower():
                                href = child.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                                if href:
                                    text += f"[Link: {run.text}]\n"
        
        # Extract tables
        if doc.tables:
            text += "\n=== TABLES ===\n"
            for table_idx, table in enumerate(doc.tables, 1):
                text += f"\nTable {table_idx}:\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text += row_text + "\n"
        
        # Extract image information
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        
        if image_count > 0:
            text += f"\n[Document contains {image_count} images]\n"
        
        return text
    except Exception as e:
        logging.error(f"Error extracting DOCX: {e}")
        return ""

def extract_text_from_image(file_bytes: bytes) -> str:
    """Extract text and image info using OCR"""
    try:
        image = Image.open(io.BytesIO(file_bytes))
        
        # Get image metadata
        text = f"[Image: {image.format}, {image.width}x{image.height}px]\n\n"
        
        # Extract text via OCR
        extracted_text = pytesseract.image_to_string(image)
        if extracted_text.strip():
            text += "=== OCR TEXT ===\n"
            text += extracted_text
        else:
            text += "[No text detected in image]"
        
        return text
    except Exception as e:
        logging.error(f"Error extracting text from image: {e}")
        return ""

async def process_document(file: UploadFile, document_id: str):
    """Process uploaded document: extract text, chunk, embed"""
    global faiss_index, chunk_metadata, metadata_index_map
    
    file_bytes = await file.read()
    file_type = file.filename.split('.')[-1].lower()
    
    # Extract text based on file type
    text = ""
    if file_type == 'pdf':
        text = extract_text_from_pdf(file_bytes)
    elif file_type in ['docx', 'doc']:
        text = extract_text_from_docx(file_bytes)
    elif file_type in ['txt', 'md']:
        raw_text = file_bytes.decode('utf-8')
        # Extract URLs from text files
        import re
        urls = re.findall(r'https?://[^\s\n]+', raw_text)
        text = raw_text
        if urls:
            text += "\n\n=== EXTRACTED LINKS ===\n"
            for url in set(urls):  # Remove duplicates
                text += f"- {url}\n"
    elif file_type in ['png', 'jpg', 'jpeg', 'bmp', 'gif']:
        text = extract_text_from_image(file_bytes)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    
    if not text.strip():
        raise HTTPException(status_code=400, detail="No text could be extracted from file")
    
    # Chunk the text
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,  # Optimized size
        chunk_overlap=150,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    chunks = text_splitter.split_text(text)
    
    # Generate embeddings in batch
    embeddings = embedding_model.encode(chunks, convert_to_numpy=True, batch_size=32)
    embeddings = embeddings.astype('float32')
    
    # Batch insert chunks into database
    chunk_docs = []
    for idx, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
        chunk_doc = DocumentChunk(
            document_id=document_id,
            chunk_index=idx,
            text=chunk,
            embedding=embedding.tolist()
        )
        chunk_dict = chunk_doc.model_dump()
        chunk_dict['upload_date'] = datetime.now(timezone.utc).isoformat()
        chunk_docs.append(chunk_dict)
    
    # Batch insert all chunks
    if chunk_docs:
        await db.document_chunks.insert_many(chunk_docs)
    
    # Update document
    await db.documents.update_one(
        {'id': document_id},
        {'$set': {'total_chunks': len(chunks), 'processed': True}}
    )
    
    # Rebuild FAISS index to ensure consistency
    await rebuild_faiss_index()
    
    return len(chunks)

async def retrieve_relevant_chunks(query: str, top_k: int = 5, document_ids: Optional[List[str]] = None) -> List[Dict]:
    """Retrieve relevant chunks using FAISS"""
    global faiss_index, chunk_metadata, metadata_index_map
    
    if faiss_index is None or len(chunk_metadata) == 0:
        return []
    
    # Embed query
    query_embedding = embedding_model.encode([query], convert_to_numpy=True).astype('float32')
    
    # Search FAISS with adjusted k
    k = min(top_k * 3, len(chunk_metadata))  # Get more for filtering
    distances, indices = faiss_index.search(query_embedding, k)
    
    # Get chunks with metadata - use dict lookup for O(1) performance
    results = []
    for idx, distance in zip(indices[0], distances[0]):
        if idx < len(metadata_index_map):
            chunk_id = metadata_index_map[idx]
            meta = chunk_metadata.get(chunk_id)
            
            if not meta:
                continue
            
            # Filter by document_ids if provided
            if document_ids and meta['document_id'] not in document_ids:
                continue
            
            # Skip if distance is too high (poor match)
            if distance > 2.0:
                continue
            
            similarity = 1.0 / (1.0 + distance)  # Convert distance to similarity
            results.append({
                'chunk_id': chunk_id,
                'document_id': meta['document_id'],
                'document_name': meta['document_name'],
                'text': meta['text'],
                'similarity': float(similarity),
                'distance': float(distance)
            })
    
    # Sort by similarity and return top_k
    results.sort(key=lambda x: x['similarity'], reverse=True)
    return results[:top_k]

def calculate_faithfulness_score(answer: str, citations: List[Citation]) -> float:
    """Calculate faithfulness score based on citation usage"""
    if not citations:
        return 0.0
    
    # Check if citations are referenced in answer
    citation_indicators = ['[1]', '[2]', '[3]', '[4]', '[5]', 'based on', 'according to']
    has_citations = any(indicator in answer.lower() for indicator in citation_indicators)
    
    # Average similarity of citations
    avg_similarity = sum(c.similarity for c in citations) / len(citations)
    
    # Combine factors
    score = (avg_similarity * 0.7) + (0.3 if has_citations else 0.0)
    return min(score, 1.0)

async def generate_answer_with_llm(query: str, context_chunks: List[Dict], mode: str) -> str:
    """Generate answer using LLM with RAG context"""
    
    # Build context from chunks
    context = "\n\n".join([
        f"[Source {idx + 1} - {chunk['document_name']}]\n{chunk['text']}"
        for idx, chunk in enumerate(context_chunks)
    ])
    
    # Calculate number of available sources
    num_sources = len(context_chunks)
    
    # Create mode-specific prompts
    mode_instructions = {
        "concise": "Provide a concise, direct answer in 2-3 sentences.",
        "detailed": "Provide a comprehensive answer with detailed explanation.",
        "research": "Provide an in-depth research-style answer with proper academic structure and extensive citations."
    }
    
    system_message = f"""You are a sophisticated AI assistant providing comprehensive, narrative-style answers based on document analysis.

CONTENT GUIDELINES:
- Answer exclusively using information from the provided documents
- If information is unavailable, state: "I don't have information about this in the provided documents."
- Write naturally without phrases like "based on the provided documents" or "the documents state"
- Present information as flowing narrative, not as rigid lists

WRITING STYLE:
- Write in a sophisticated, professional narrative format
- Use flowing paragraphs that connect ideas naturally
- Integrate information smoothly rather than listing points
- Use bullet points ONLY when comparing specific items or when a list genuinely improves clarity
- Avoid over-structured formats (e.g., "What is it?", "What does it do?")
- Create cohesive, essay-like responses that read naturally
- Use transitions between ideas (however, furthermore, additionally, consequently)

FORMATTING (Markdown):
- Use **bold** for key terms and important concepts
- Use *italic* for emphasis and technical terms on first mention
- Use headings (#, ##) sparingly - only for major topic shifts
- Prefer paragraph-based explanations over bullet lists
- Use tables only for genuine data comparison
- Keep formatting subtle and professional

CITATIONS (CRITICAL):
- You have access to exactly {num_sources} sources: [1] through [{num_sources}]
- NEVER use citation numbers higher than [{num_sources}]
- Place citations immediately after the relevant statement: "The system processes data efficiently [1]"
- Integrate citations naturally into flowing text
- You may cite the same source multiple times
- Do NOT group citations at the end of paragraphs
- Example: "The platform achieves 95% accuracy [1] while maintaining low latency [2], making it suitable for real-time applications [1]"

ANSWER STRUCTURE:
- Begin with a clear, direct answer to the question
- Develop the answer with supporting details in subsequent paragraphs
- Integrate examples and specifics naturally
- Conclude with implications or significance when relevant
- Maintain narrative flow throughout

Output sophisticated, well-written prose with inline citations [1] to [{num_sources}] embedded naturally in the text."""
    
    user_prompt = f"""Context Documents:
{context}

Question: {query}

Answer:"""
    
    try:
        # Initialize Azure OpenAI client
        client = AzureOpenAI(
            api_key=os.environ.get('AZURE_OPENAI_API_KEY'),
            api_version="2024-08-01-preview",
            azure_endpoint=os.environ.get('AZURE_OPENAI_ENDPOINT')
        )
        
        # Generate answer
        response = client.chat.completions.create(
            model=os.environ.get('AZURE_OPENAI_DEPLOYMENT', 'gpt-5.2-chat'),
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_prompt}
            ]
        )
        
        return response.choices[0].message.content
    except Exception as e:
        logging.error(f"LLM error: {e}")
        raise HTTPException(status_code=500, detail="Error generating answer")

# API Routes
@api_router.get("/")
async def root():
    return {"message": "NeuroQuery RAG API"}

@api_router.post("/documents/upload", response_model=Document)
async def upload_document(file: UploadFile = File(...)):
    """Upload and process a document"""
    
    # Create document record
    doc = Document(
        filename=file.filename,
        file_type=file.filename.split('.')[-1].lower()
    )
    
    doc_dict = doc.model_dump()
    doc_dict['upload_date'] = doc_dict['upload_date'].isoformat()
    await db.documents.insert_one(doc_dict)
    
    # Process document in background
    try:
        chunks_count = await process_document(file, doc.id)
        logging.info(f"Processed document {file.filename}: {chunks_count} chunks")
    except Exception as e:
        logging.error(f"Error processing document: {e}")
        await db.documents.delete_one({'id': doc.id})
        raise HTTPException(status_code=500, detail=str(e))
    
    return doc

@api_router.get("/documents", response_model=List[Document])
async def get_documents(skip: int = 0, limit: int = 100):
    """Get all uploaded documents with pagination"""
    docs = await db.documents.find({}, {"_id": 0}).skip(skip).limit(limit).to_list(None)
    for doc in docs:
        if isinstance(doc.get('upload_date'), str):
            doc['upload_date'] = datetime.fromisoformat(doc['upload_date'])
    return docs

@api_router.delete("/documents/{document_id}")
async def delete_document(document_id: str):
    """Delete a document and its chunks"""
    global chunk_metadata, metadata_index_map
    
    # Delete from database
    await db.documents.delete_one({'id': document_id})
    await db.document_chunks.delete_many({'document_id': document_id})
    
    # Remove from metadata - filter dict keys
    chunk_ids_to_remove = [cid for cid, meta in chunk_metadata.items() 
                          if meta['document_id'] == document_id]
    for chunk_id in chunk_ids_to_remove:
        del chunk_metadata[chunk_id]
    
    # Update metadata_index_map
    metadata_index_map = [cid for cid in metadata_index_map if cid not in chunk_ids_to_remove]
    
    return {"message": "Document deleted"}

@api_router.post("/query", response_model=QueryResponse)
async def query_documents(request: QueryRequest):
    """Query documents using RAG"""
    
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")
    
    # Retrieve relevant chunks
    top_k = 5 if request.mode == "concise" else 8
    chunks = await retrieve_relevant_chunks(
        request.query,
        top_k=top_k,
        document_ids=request.document_ids
    )
    
    if not chunks:
        return QueryResponse(
            answer="No relevant information found in the uploaded documents.",
            citations=[],
            faithfulness_score=0.0,
            refused=True
        )
    
    # Generate answer
    answer = await generate_answer_with_llm(request.query, chunks, request.mode)
    
    # Check if LLM refused to answer
    refused = "cannot answer" in answer.lower() or "don't have" in answer.lower()
    
    # Create citations
    citations = [
        Citation(
            chunk_id=chunk['chunk_id'],
            document_id=chunk['document_id'],
            document_name=chunk['document_name'],
            text=chunk['text'][:300] + "..." if len(chunk['text']) > 300 else chunk['text'],
            similarity=chunk['similarity']
        )
        for chunk in chunks
    ]
    
    # Calculate faithfulness score
    faithfulness_score = calculate_faithfulness_score(answer, citations)
    if refused:
        faithfulness_score = 0.0
    
    # Retrieval details for debug panel
    retrieval_details = {
        'retrieved_chunks': len(chunks),
        'chunks': [
            {
                'document': chunk['document_name'],
                'similarity': chunk['similarity'],
                'distance': chunk['distance']
            }
            for chunk in chunks
        ]
    }
    
    return QueryResponse(
        answer=answer,
        citations=citations,
        faithfulness_score=faithfulness_score,
        retrieval_details=retrieval_details,
        refused=refused
    )

@api_router.get("/documents/{document_id}/chunks")
async def get_document_chunks(document_id: str, skip: int = 0, limit: int = 100):
    """Get all chunks for a document with pagination"""
    chunks = await db.document_chunks.find(
        {'document_id': document_id},
        {'_id': 0, 'embedding': 0}  # Exclude embeddings
    ).skip(skip).limit(limit).to_list(None)
    return chunks

# Chat Routes
@api_router.post("/chats", response_model=ChatSession)
async def create_chat(document_ids: Optional[List[str]] = None):
    """Create a new chat session"""
    chat = ChatSession(
        document_ids=document_ids or []
    )
    chat_dict = chat.model_dump()
    chat_dict['created_at'] = chat_dict['created_at'].isoformat()
    chat_dict['updated_at'] = chat_dict['updated_at'].isoformat()
    await db.chat_sessions.insert_one(chat_dict)
    return chat

@api_router.get("/chats/{chat_id}", response_model=ChatSession)
async def get_chat(chat_id: str):
    """Get a chat session with all messages"""
    chat = await db.chat_sessions.find_one({'id': chat_id})
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    return chat

@api_router.post("/chats/{chat_id}/messages")
async def send_message(chat_id: str, message: str, mode: str = "detailed"):
    """Send a message in a chat session and get AI response"""
    
    # Get chat session
    chat = await db.chat_sessions.find_one({'id': chat_id})
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    
    if not message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")
    
    # Add user message
    user_msg = ChatMessage(role="user", content=message)
    user_msg_dict = user_msg.model_dump()
    user_msg_dict['timestamp'] = user_msg_dict['timestamp'].isoformat()
    
    # Retrieve relevant chunks
    document_ids = chat.get('document_ids') or None
    top_k = 5 if mode == "concise" else 8
    chunks = await retrieve_relevant_chunks(
        message,
        top_k=top_k,
        document_ids=document_ids
    )
    
    if not chunks:
        answer = "I don't have any relevant information in the uploaded documents to answer your question. Could you please upload relevant documents first?"
        citations = []
    else:
        # Generate answer
        answer = await generate_answer_with_llm(message, chunks, mode)
        
        # Create citations
        citations = [
            {
                "chunk_id": chunk['chunk_id'],
                "document_id": chunk['document_id'],
                "document_name": chunk['document_name'],
                "text": chunk['text'][:300] + "..." if len(chunk['text']) > 300 else chunk['text'],
                "similarity": chunk['similarity']
            }
            for chunk in chunks
        ]
    
    # Add assistant message
    assistant_msg = ChatMessage(role="assistant", content=answer, citations=citations)
    assistant_msg_dict = assistant_msg.model_dump()
    assistant_msg_dict['timestamp'] = assistant_msg_dict['timestamp'].isoformat()
    
    # Update chat with both messages
    messages = chat.get('messages', [])
    messages.append(user_msg_dict)
    messages.append(assistant_msg_dict)
    
    # Update title if first message
    title = chat.get('title', 'New Chat')
    if len(messages) == 2:  # First exchange
        title = message[:50] + "..." if len(message) > 50 else message
    
    await db.chat_sessions.update_one(
        {'id': chat_id},
        {'$set': {
            'messages': messages,
            'title': title,
            'updated_at': datetime.now(timezone.utc).isoformat()
        }}
    )
    
    return {
        "user_message": user_msg,
        "assistant_message": assistant_msg,
        "citations": citations
    }

@api_router.get("/chats")
async def list_chats():
    """List all chat sessions"""
    chats = await db.chat_sessions.find({}, {"_id": 0}).sort('updated_at', -1).to_list(100)
    return chats

@api_router.delete("/chats/{chat_id}")
async def delete_chat(chat_id: str):
    """Delete a chat session"""
    await db.chat_sessions.delete_one({'id': chat_id})
    return {"message": "Chat deleted"}

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=os.environ.get('CORS_ORIGINS', '*').split(','),
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    client.close()

@app.on_event("startup")
async def startup_db_client():
    """Rebuild FAISS index from MongoDB on startup"""
    await rebuild_faiss_index()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)